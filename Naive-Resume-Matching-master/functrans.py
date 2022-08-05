
import os
import nltk
#nltk.download('punkt')
#nltk.download('stopwords')
import io
import docx2txt
from docx import Document
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
import pandas as pd
import textwrap
from deep_translator import GoogleTranslator
from pyresparser import ResumeParser
import os.path  

import os,requests
from pyresparser import ResumeParser
import json
import pandas as pd
from ast import literal_eval
import re
import urllib.request
import urllib.parse
import requests
from io import StringIO, BytesIO


#####################################################################################################################
############# Fonction de traduction interne


# Fonction list to string
def listToString(s): 
    str1 = " "  
    return (str1.join(s))

#lecture et traduction pdf
def read_pdf(path):
    rsrcmgr = PDFResourceManager()
    retstr = io.StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    fp = open(path, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos=set()
    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password, caching=caching, check_extractable=True): 
        interpreter.process_page(page)
    text = retstr.getvalue()
    text = " ".join(text.replace(u"\xa0", " ").strip().split())
    lines = textwrap.wrap(text, 4999, break_long_words=False)
    textlist=[]
    for line in lines:
        transtext=line
        #text= GoogleTranslator(source='auto', target='en').translate(transtext)
        #text=mytrans(transtext)
        textlist.append(text)
    text=listToString(textlist)    
    fp.close()
    device.close()
    retstr.close()
    return text

#lecture et traduction docx doc
def docx2txtrans(dirfile):
    text=docx2txt.process(dirfile)
    text = " ".join(text.replace(u"\xa0", " ").strip().split())
    #lines = textwrap.wrap(text, 4999, break_long_words=False)
    lines = textwrap.wrap(text, 4999, break_long_words=False)
    textlist=[]
    for line in lines:
        transtext=line
        #text= GoogleTranslator(source='auto', target='en').translate(transtext)
        #text=mytrans(transtext)
        textlist.append(text)
    text=listToString(textlist)    
    return text

#détecteur de format
def read_resumes(list_of_resumes, resume_directory):
    placeholder = []
    for res in list_of_resumes:
        temp = []
        temp.append(res)
        dirfile=resume_directory+res
        if (dirfile.endswith('.docx')):
            text = docx2txtrans(dirfile)
        elif (dirfile.endswith('.doc')):
            text = docx2txtrans(dirfile)
        elif (dirfile.endswith('.pdf')):
            text = read_pdf(dirfile)
        temp.append(text)
        placeholder.append(temp)
    return placeholder


#traducteur text venant de dataframe
def translator(lines):
    textlist=[]
    for line in lines:
        transtext=line
        #text= GoogleTranslator(source='auto', target='en').translate(transtext)
        #text=mytrans(transtext)
        textlist.append(text)
    text=listToString(textlist)    
    return text


# fonction d'extraction des compétence d'un CV
def skill_resumes(list_of_resumes, resume_directory):
    placeholder = []
    for res in list_of_resumes:
        temp = []
        temp.append(res)
        dirfile=resume_directory+res
        if (dirfile.endswith('.docx')):
            text = ResumeParser(dirfile).get_extracted_data()
        elif (dirfile.endswith('.doc')):
            text = ResumeParser(dirfile).get_extracted_data()
        elif (dirfile.endswith('.pdf')):
            text = ResumeParser(dirfile).get_extracted_data()
        temp.append(text)
        placeholder.append(temp)
    return placeholder


#fonction de récupération  des skill d'un job description

def skills_doc(text):
    ress = []
    skill_job=[]
    document = Document()
    document.add_paragraph(text)
    document.save('docx_file.docx')
    text = ResumeParser('docx_file.docx').get_extracted_data()
    skill_job.append(text)
    os.remove("docx_file.docx")
    [ress.append(x) for x in skill_job if x not in ress]
    return ress
    

# itération pour recupéré les skills et le nom du job et resum
def get_skjob_resum(dict_job_resum):
    L="skills"
    M="name"
    sklst=[]
    sklst2=[]
    nmlst=[]
    for i in dict_job_resum:
        l=i
        res = None
        resn=None
        if all(L in sub for sub in [l]):
            res = l[L]
        if all(M in sub for sub in [l]):
            resn = l[M]

        sklst.append(res)
        nmlst.append(resn)
        

    #Creation de la dataframe Nom et skill

    df = pd.DataFrame({'Nom': nmlst,'Skills': sklst})

    def listToStringWithoutBrackets(list1):
        return str(list1).replace('[','').replace(']','').replace("'",'')

    df['Skills']=df['Skills'].apply(listToStringWithoutBrackets)
    return df

    ##################################################################################""
    #extration de text depuis des url pdf
def extract_text_from_pdf_url(url, user_agent=None):
    resource_manager = PDFResourceManager()
    fake_file_handle = StringIO()
    converter = TextConverter(resource_manager, fake_file_handle)    

    if user_agent == None:
        user_agent = 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/83.0.4103.61 Safari/537.36'

    headers = {'User-Agent': user_agent}
    request = urllib.request.Request(url, data=None, headers=headers)

    response = urllib.request.urlopen(request).read()
    fb = BytesIO(response)

    page_interpreter = PDFPageInterpreter(resource_manager, converter)

    for page in PDFPage.get_pages(fb,
                                caching=True,
                                check_extractable=True):
        page_interpreter.process_page(page)


    text = fake_file_handle.getvalue()

    # close open handles
    fb.close()
    converter.close()   
    fake_file_handle.close()

    if text:
        # If document has instances of \xa0 replace them with spaces.
        # NOTE: \xa0 is non-breaking space in Latin1 (ISO 8859-1) & chr(160)
        text = text.replace(u'\xa0', u' ')
        lines = textwrap.wrap(text, 4999, break_long_words=False)
        textlist=[]
        for line in lines:
            transtext=line
            #text= GoogleTranslator(source='auto', target='en').translate(transtext)
            #text=mytrans(transtext)
            textlist.append(text)
        text=listToString(textlist)

        return text
    

#extration de text depuis des url docx    
def extract_text_from_docx_url(url, user_agent=None):
    
    if user_agent == None:
        user_agent = 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/83.0.4103.61 Safari/537.36'

    headers = {'User-Agent': user_agent}
    request = urllib.request.Request(url, data=None, headers=headers)

    response = urllib.request.urlopen(request).read()
    fb = BytesIO(response)
    text = docx2txt.process(fb)
    lines = textwrap.wrap(text, 4999, break_long_words=False)
    textlist=[]
    for line in lines:
        transtext=line
        #text= GoogleTranslator(source='auto', target='en').translate(transtext)
        #text=mytrans(transtext)
        textlist.append(text)
    text=listToString(textlist)
    return text




def skills_doc2(text):
    ress = []
    skill_c=[]
    document = Document()
    document.add_paragraph(text)
    document.save('docx_file.docx')
    text = ResumeParser('docx_file.docx').get_extracted_data()
    skill_c.append(text)
    os.remove("docx_file.docx")
    [ress.append(x) for x in skill_c if x not in ress]
    return ress


    

""" 
def download(url):
    get_response = requests.get(url,stream=True)
    file_name  = url.split("/")[-1]
    with open(file_name, 'wb') as f:
        for chunk in get_response.iter_content(chunk_size=1024):
            if chunk: # filter out keep-alive new chunks
                f.write(chunk)

"""

def download(url):
    get_response = requests.get(url,stream=True)
    file_name  = url.split("/")[-1]
    directory = './CVdir/'
    file_path = os.path.join(directory, file_name)
    if not os.path.isdir(directory):
        os.mkdir(directory)
    with open(file_path, 'wb') as f:
        for chunk in get_response.iter_content(chunk_size=1024):
            if chunk: # filter out keep-alive new chunks
                f.write(chunk)
                
